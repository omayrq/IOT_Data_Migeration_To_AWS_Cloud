import os
from docx import Document
from docx.shared import Inches

# Paths
project_root = r"E:/Hackathon/Hackathon/iot-streaming-platform/iot-streaming-platform"
image_path = r"C:/Users/MYC/.gemini/antigravity/brain/4ff72b83-b249-4041-8267-98e42023b255/iot_platform_architecture_1783854855845.jpg"
output_path = os.path.join(project_root, "docs", "architecture.docx")

# Create document
doc = Document()
doc.add_heading('IoT Streaming Platform Architecture', level=1)
doc.add_paragraph('The diagram below illustrates the production‑ready architecture of the platform.')

if os.path.exists(image_path):
    doc.add_picture(image_path, width=Inches(6))
else:
    doc.add_paragraph('**Image not found at expected path.**')

# Save
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Document created at {output_path}")
